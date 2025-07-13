import streamlit as st
import os
import pandas as pd
import tempfile
import asyncio
from concurrent.futures import ThreadPoolExecutor
from langchain.document_loaders import PyPDFLoader
from langchain.embeddings import HuggingFaceEmbeddings
from langchain import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.schema import Document
import docx
import speech_recognition as sr
from gtts import gTTS
import io
import pygame
from transformers import pipeline

# Set Google API Key
os.environ["GOOGLE_API_KEY"] = "AIzaSyAwZeFnJNdT7cc7ze4MMyCD8ZX9a_66WTw"  # Replace with your API key

# Initialize summarization pipeline
@st.cache_resource
def get_summarizer():
    return pipeline("summarization", model="facebook/bart-large-cnn")

# Create a document database from the uploaded file
def create_document_db(file_path, file_type):
    if file_type == 'pdf':
        loader = PyPDFLoader(file_path)
        pages = loader.load_and_split()
    elif file_type == 'csv':
        df = pd.read_csv(file_path)
        pages = [Document(page_content=" | ".join([f"{col}: {val}" for col, val in row.items()])) for _, row in df.iterrows()]
    elif file_type == 'txt':
        with open(file_path, 'r') as file:
            text = file.read()
        pages = [Document(page_content=text)]
    elif file_type == 'docx':
        doc = docx.Document(file_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        pages = [Document(page_content=text)]
    else:
        raise ValueError("Unsupported file type")

    embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    db = FAISS.from_documents(pages, embeddings)
    return db

# Handle user query and return answer
def handle_query(file_path, file_type, query):
    db = create_document_db(file_path, file_type)
    docs = db.similarity_search(query)
    content = "\n".join([doc.page_content for doc in docs])
    
    qa_prompt = ("Use the following pieces of context to answer the user's question. "
                 "If you don't know the answer, just say that you don't know, don't try to make up an answer."
                 "----------------")
    input_text = qa_prompt + "\nContext:" + content + "\nUser question:\n" + query

    llm = ChatGoogleGenerativeAI(model="gemini-pro")
    result = llm.invoke(input_text)
    
    return result.content

# Recognize speech input from microphone
async def recognize_speech():
    recognizer = sr.Recognizer()
    with sr.Microphone() as source:
        st.info("Listening for your query...")
        audio = recognizer.listen(source, timeout=5, phrase_time_limit=10)
        try:
            text = recognizer.recognize_google(audio)
            st.success(f"You said: {text}")
            return text
        except sr.UnknownValueError:
            st.error("Could not understand the audio. Please try again.")
            return None
        except sr.RequestError:
            st.error("Error with the speech recognition request. Please check your connection.")
            return None

# Convert text to speech
def speak_text(text):
    tts = gTTS(text=text, lang='en')
    fp = io.BytesIO()
    tts.write_to_fp(fp)
    fp.seek(0)
    pygame.mixer.init()
    pygame.mixer.music.load(fp)
    pygame.mixer.music.play()
    while pygame.mixer.music.get_busy():
        pygame.time.Clock().tick(10)

# Summarize the document
def summarize_document(file_path, file_type):
    summarizer = get_summarizer()
    db = create_document_db(file_path, file_type)
    content = "\n".join([doc.page_content for doc in db.similarity_search("")])
    
    summary = summarizer(content, max_length=150, min_length=50, do_sample=False)
    return summary[0]['summary_text']

# Custom CSS for UI
st.markdown("""
    <style>
        .stButton > button {
            background-color: #4CAF50;
            color: white;
            border-radius: 8px;
            padding: 10px 20px;
            font-size: 16px;
        }
        .stTextInput > div > input {
            border: 2px solid #4CAF50;
            border-radius: 8px;
            padding: 8px;
            font-size: 16px;
        }
        .stRadio > div > label {
            font-size: 16px;
            font-weight: bold;
        }
        .stFileUploader > label {
            font-size: 16px;
            font-weight: bold;
            color: #4CAF50;
        }
        .stSpinner > div > div > div {
            color: #4CAF50;
        }
        .stInfo {
            background-color: #f1f8e9;
            border-radius: 8px;
            padding: 20px;
            font-size: 16px;
        }
        .stSuccess {
            background-color: #e8f5e9;
            border-radius: 8px;
            padding: 10px;
            font-size: 16px;
            color: #388e3c;
        }
        .stError {
            background-color: #ffebee;
            border-radius: 8px;
            padding: 10px;
            font-size: 16px;
            color: #d32f2f;
        }
        .main {
            background-color: #f4f7f6;
            padding: 20px;
            border-radius: 10px;
            margin: 20px;
        }
    </style>
    """, unsafe_allow_html=True)

# Streamlit interface
st.title("📄 Document Interaction Platform")
st.write("Upload a document (PDF, CSV, DOCX, TXT) and interact through text or voice queries. You can also request a document summary.")

# File upload
uploaded_file = st.file_uploader("📁 Choose a file", type=['pdf', 'csv', 'txt', 'docx'])
file_type = None

if uploaded_file is not None:
    if uploaded_file.name.endswith('.pdf'):
        file_type = 'pdf'
    elif uploaded_file.name.endswith('.csv'):
        file_type = 'csv'
    elif uploaded_file.name.endswith('.txt'):
        file_type = 'txt'
    elif uploaded_file.name.endswith('.docx'):
        file_type = 'docx'
    
    file_path = os.path.join(tempfile.gettempdir(), uploaded_file.name)
    with open(file_path, 'wb') as f:
        f.write(uploaded_file.getbuffer())
    
    st.success(f"📄 Uploaded file: {uploaded_file.name}")
    
    # Summarize Document
    if st.button("📋 Summarize Document"):
        with st.spinner('Summarizing...'):
            summary = summarize_document(file_path=file_path, file_type=file_type)
        st.write("📝 **Summary:**")
        st.write(summary)
    
    # Interaction Options
    st.write("### 🛠️ How would you like to interact?")
    interaction_mode = st.radio("Select an option:", ("📝 Text Query", "🎤 Voice Query"))
    
    if interaction_mode == "📝 Text Query":
        query = st.text_input("🔍 Enter your query:")
        if st.button("📩 Submit Query"):
            if query:
                with st.spinner('Processing your query...'):
                    result = handle_query(file_path=file_path, file_type=file_type, query=query)
                st.write("🔍 **Answer:**")
                st.write(result)
                if st.button("🔊 Speak Answer"):
                    speak_text(result)
            else:
                st.warning("⚠️ Please enter a query.")
    
    elif interaction_mode == "🎤 Voice Query":
        if st.button("🎙️ Record Voice Query"):
            with st.spinner('Listening...'):
                query = asyncio.run(recognize_speech())
            if query:
                with st.spinner('Processing your query...'):
                    result = handle_query(file_path=file_path, file_type=file_type, query=query)
                st.write("🔍 **Answer:**")
                st.write(result)
                if st.button("🔊 Speak Answer"):
                    speak_text(result)

# Provide guidance and feedback
st.info("🔑 Ensure your Google API key is set and valid. For voice interaction, please allow microphone access.")
